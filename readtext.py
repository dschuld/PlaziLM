import docx
import os

DOC_FOLDER = os.environ["PLAZI_DOCS_FOLDER"]


def read_text_file(filename):
    doc = docx.Document(DOC_FOLDER + filename)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text

    return text

def read_text(filename, num_words):
    doc = docx.Document(DOC_FOLDER + filename)

    word_count = 0
    text = ''
    for paragraph in doc.paragraphs:
        if word_count >= num_words:
            break
        text += paragraph.text
        word_count += len(paragraph.text.split())

    return text


def read_text_from_offset(filename, start_offset, num_words):
    doc = docx.Document(DOC_FOLDER + filename)

    word_count = 0
    text = ''
    for paragraph in doc.paragraphs:
        paragraph_words = paragraph.text.split()
        paragraph_word_count = len(paragraph_words)
        if word_count + paragraph_word_count < start_offset:
            word_count += paragraph_word_count
            continue
        if word_count >= start_offset + num_words:
            break
        text += paragraph.text + ' '
        word_count += paragraph_word_count

    return text

def write_to_file(filename, text):
    doc = docx.Document()
    doc.add_paragraph(text)
    doc.add_paragraph('\n\n\n\n\n\n')
    doc.save(DOC_FOLDER + filename)



